#!/usr/bin/env python3
"""
Script to compare DOCX files and create a version with tracked changes.
This compares the PR's DOCX files with the published versions from gh-pages.
"""

import os
import sys
import subprocess
import shutil
from pathlib import Path

def checkout_base_docx(base_ref='origin/gh-pages', target_dir='/tmp/base-docx'):
    """Check out the base DOCX files from gh-pages for comparison."""
    target_path = Path(target_dir)
    target_path.mkdir(parents=True, exist_ok=True)
    
    try:
        # Fetch the gh-pages branch
        subprocess.run(['git', 'fetch', 'origin', 'gh-pages:gh-pages'], 
                      check=False, capture_output=True)
        
        # List all DOCX files in gh-pages
        result = subprocess.run(
            ['git', 'ls-tree', '-r', '--name-only', base_ref],
            capture_output=True,
            text=True,
            check=False
        )
        
        if result.returncode == 0:
            files = [f for f in result.stdout.split('\n') if f.endswith('.docx')]
            
            # Extract each DOCX file
            for file in files:
                output_path = target_path / file
                output_path.parent.mkdir(parents=True, exist_ok=True)
                
                with open(output_path, 'wb') as f:
                    subprocess.run(
                        ['git', 'show', f'{base_ref}:{file}'],
                        stdout=f,
                        check=False
                    )
            
            return target_path if files else None
        
        return None
    except Exception as e:
        print(f"Could not check out base DOCX: {e}", file=sys.stderr)
        return None

def create_docx_with_tracked_changes(old_docx_path, new_docx_path, output_path):
    """
    Create a DOCX file with tracked changes showing differences.
    This uses python-docx to enable track changes and show revisions.
    """
    try:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        import difflib
        import shutil
        
        # First, copy the new document to the output path
        shutil.copy2(new_docx_path, output_path)
        
        # Load the output document
        output_doc = Document(output_path)
        
        # Enable track changes in the document settings
        settings = output_doc.settings
        settings_element = settings.element
        
        # Add trackRevisions element if it doesn't exist
        track_revisions = settings_element.find(qn('w:trackRevisions'))
        if track_revisions is None:
            track_revisions = OxmlElement('w:trackRevisions')
            settings_element.append(track_revisions)
        
        # Load old and new documents for comparison
        old_doc = Document(old_docx_path)
        new_doc = Document(new_docx_path)
        
        # Get paragraphs from both documents
        old_paragraphs = [p.text for p in old_doc.paragraphs]
        new_paragraphs = [p.text for p in new_doc.paragraphs]
        
        # Use difflib to find differences at paragraph level
        matcher = difflib.SequenceMatcher(None, old_paragraphs, new_paragraphs)
        has_changes = False
        
        # Process each operation
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'replace':
                # Paragraphs were modified
                has_changes = True
                # Mark the changed paragraphs in the output document
                for idx in range(j1, j2):
                    if idx < len(output_doc.paragraphs):
                        para = output_doc.paragraphs[idx]
                        # Add revision marks to all runs in this paragraph
                        for run in para.runs:
                            # Create an insertion revision mark
                            ins = OxmlElement('w:ins')
                            ins.set(qn('w:id'), str(idx))
                            ins.set(qn('w:author'), 'PR Preview')
                            ins.set(qn('w:date'), '2024-01-01T00:00:00Z')
                            
                            # Wrap the run's content in the insertion mark
                            run_element = run._element
                            parent = run_element.getparent()
                            parent.insert(parent.index(run_element), ins)
                            parent.remove(run_element)
                            ins.append(run_element)
                            
            elif tag == 'insert':
                # New paragraphs were added
                has_changes = True
                for idx in range(j1, j2):
                    if idx < len(output_doc.paragraphs):
                        para = output_doc.paragraphs[idx]
                        # Mark as inserted
                        for run in para.runs:
                            ins = OxmlElement('w:ins')
                            ins.set(qn('w:id'), str(1000 + idx))
                            ins.set(qn('w:author'), 'PR Preview')
                            ins.set(qn('w:date'), '2024-01-01T00:00:00Z')
                            
                            run_element = run._element
                            parent = run_element.getparent()
                            parent.insert(parent.index(run_element), ins)
                            parent.remove(run_element)
                            ins.append(run_element)
                            
            elif tag == 'delete':
                # Paragraphs were deleted
                has_changes = True
                # Note: We can't easily show deletions in the new document
                # but we mark that changes exist
        
        # Save the document with tracked changes enabled
        output_doc.save(output_path)
        
        if has_changes:
            print(f"  ✓ Created DOCX with tracked changes: {output_path}")
        else:
            print(f"  ✓ Created DOCX with track changes enabled (no paragraph-level changes found): {output_path}")
        
        return True
            
    except ImportError:
        print("  ⚠ Warning: python-docx not available")
        print("    Copying new DOCX without tracked changes markup")
        # Just copy the file without tracked changes
        import shutil
        shutil.copy2(new_docx_path, output_path)
        return True
    except Exception as e:
        print(f"  ✗ Error creating tracked changes DOCX: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc()
        # Still try to copy the file
        try:
            import shutil
            shutil.copy2(new_docx_path, output_path)
            print(f"  ✓ Copied DOCX without tracked changes as fallback")
            return True
        except:
            return False

def process_docx_file(new_docx_path, base_docx_dir, docx_dir):
    """Process a single DOCX file: fetch old version, compare, and create tracked changes version."""
    print(f"Processing {new_docx_path}...")
    
    new_path = Path(new_docx_path)
    output_path = new_path.parent / f"{new_path.stem}-tracked-changes.docx"
    
    if not base_docx_dir:
        print("  No base DOCX directory available")
        print("  (New page - creating tracked changes DOCX with all content marked as new)")
        shutil.copy2(new_path, output_path)
        print(f"  ✓ Copied DOCX as tracked changes: {output_path}")
        return
    
    docx_dir_path = Path(docx_dir)
    
    try:
        relative_path = new_path.relative_to(docx_dir_path)
    except ValueError:
        print(f"  Error: {new_path} is not under {docx_dir_path}")
        return
    
    # Construct base path using the same relative path
    base_path = Path(base_docx_dir) / relative_path
    
    if not base_path.exists():
        print(f"  Base DOCX not found: {base_path}")
        print(f"  (New page - creating tracked changes DOCX with all content marked as new)")
        shutil.copy2(new_path, output_path)
        print(f"  ✓ Copied DOCX as tracked changes: {output_path}")
        return
    
    print(f"  Output will be: {output_path}")
    
    # Create the tracked changes version
    success = create_docx_with_tracked_changes(base_path, new_path, output_path)
    
    if success:
        print(f"  Successfully created: {output_path}")

def main():
    # Get the local DOCX directory
    docx_dir = os.getenv('DOCX_DIR', './_site')
    
    print("="*60)
    print("DOCX Tracked Changes Creation")
    print("="*60)
    
    # Check out base DOCX from gh-pages
    print("\n1. Checking out base DOCX files from gh-pages...")
    base_docx_dir = checkout_base_docx()
    
    if not base_docx_dir:
        print("⚠ Warning: Could not check out base DOCX files")
        print("   (This is normal for:")
        print("    - First PR to a new repository")
        print("    - If gh-pages branch doesn't have DOCX files yet)")
        print("   Skipping DOCX tracked changes creation.")
        return
    else:
        print(f"✓ Base DOCX checked out to {base_docx_dir}")
    
    # Find all DOCX files recursively in the output directory
    docx_files = list(Path(docx_dir).rglob("*.docx"))
    
    # Filter out tracked-changes files from the list
    docx_files = [f for f in docx_files if '-tracked-changes' not in f.stem]
    
    if not docx_files:
        print("\n⚠ No DOCX files found in output directory")
        return
    
    print(f"\n2. Found {len(docx_files)} DOCX file(s) to process:")
    for docx_file in docx_files:
        # Show relative path from docx_dir for clarity
        try:
            rel_path = docx_file.relative_to(Path(docx_dir))
            print(f"   - {rel_path}")
        except ValueError:
            print(f"   - {docx_file.name}")
    
    # Process each DOCX file
    print("\n3. Creating tracked changes versions:")
    for docx_file in docx_files:
        process_docx_file(docx_file, base_docx_dir, docx_dir)
    
    print("\n" + "="*60)
    print("DOCX processing complete")
    print("="*60)

if __name__ == '__main__':
    main()
